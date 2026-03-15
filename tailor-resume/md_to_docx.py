"""
Markdown to Word converter for tailor-resume skill.
Uses python-docx to create Word documents from Markdown content.
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


def set_chinese_font(run, font_name="微软雅黑", font_size=11):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def parse_markdown_to_word(md_content: str, output_path: str, style_reference: str = None):
    doc = Document()
    
    if style_reference and Path(style_reference).exists():
        try:
            from docx import Document as DocLoader
            style_doc = DocLoader(style_reference)
            doc.styles = style_doc.styles
        except Exception:
            pass
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            for run in p.runs:
                set_chinese_font(run, "微软雅黑", 22)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=1)
            for run in p.runs:
                set_chinese_font(run, "微软雅黑", 16)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=2)
            for run in p.runs:
                set_chinese_font(run, "微软雅黑", 14)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            run = p.add_run(text)
            set_chinese_font(run, "微软雅黑", 11)
        elif line.startswith('  - ') or line.startswith('  * '):
            p = doc.add_paragraph(style='List Bullet 2')
            text = line[4:]
            run = p.add_run(text)
            set_chinese_font(run, "微软雅黑", 10)
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s', '', line)
            run = p.add_run(text)
            set_chinese_font(run, "微软雅黑", 11)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            set_chinese_font(run, "微软雅黑", 11)
        elif line == '---':
            doc.add_paragraph('_' * 50)
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, "微软雅黑", 11)
        
        i += 1
    
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: python md_to_docx.py <input.md> <output.docx> [style_reference.docx]")
        sys.exit(1)
    
    input_md = sys.argv[1]
    output_docx = sys.argv[2]
    style_ref = sys.argv[3] if len(sys.argv) > 3 else None
    
    with open(input_md, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    parse_markdown_to_word(md_content, output_docx, style_ref)


if __name__ == "__main__":
    main()
